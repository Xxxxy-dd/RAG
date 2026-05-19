"""DOCX loader to Markdown format.

专属格式深度解析，将docx统一输出为 Markdown 文档，方便后续做统一切分。

改进点与设计决策：
- 优先识别 Word 自带标题、中文标题和编号标题，并转成 Markdown 标题。
- 保留正文段落、项目符号和编号列表的语义，输出为 Markdown 列表或普通段落。
- 单独抽取表格并保留单元格语义，转成 Markdown 表格，避免重要信息丢失。
- 维持章节路径和页内顺序，便于后续统一切分和检索溯源。

使用说明：优先使用 `load_docx(path)` 作为统一入口，它会返回已经 Markdown 化的 `Document` 列表；
后续统一切分工作应交给上层的 Markdown splitter。
"""

import re
from typing import List, Optional, Tuple

from docx import Document as DocxDocument
from docx.document import Document as _DocxDocumentType
from docx.table import Table
from docx.text.paragraph import Paragraph
from langchain_core.documents import Document

try:
    import tiktoken
except ImportError:  # pragma: no cover - optional dependency
    tiktoken = None


_TOKEN_ENCODER = None


def _normalize_text(text: str) -> str:
    """把乱掉的空格和换行整理干净，方便后面统一转 Markdown。"""
    if not text:
        return ""

    text = text.replace("\r\n", "\n").replace("\r", "\n")
    text = re.sub(r"[ \t\f\v]+", " ", text)
    text = re.sub(r"\n{3,}", "\n\n", text)
    return text.strip()


def _token_length(text: str) -> int:
    """粗略算一下 token 数，优先用 tiktoken，没有就用字符数近似。"""
    if not text:
        return 0

    if tiktoken is not None:
        try:
            global _TOKEN_ENCODER
            if _TOKEN_ENCODER is None:
                _TOKEN_ENCODER = tiktoken.get_encoding("cl100k_base")
            return len(_TOKEN_ENCODER.encode(text))
        except Exception:  # pragma: no cover - defensive fallback
            pass

    return max(1, len(text) // 2)


def _detect_content_type(text: str) -> str:
    """简单判断文本类型，方便后面知道它更像正文还是列表。"""
    cleaned = _normalize_text(text)
    if not cleaned:
        return "paragraph"

    lines = [line.strip() for line in cleaned.split("\n") if line.strip()]
    if not lines:
        return "paragraph"

    list_like_count = 0
    for line in lines:
        if re.match(
            r"^([\-\*•·]|\d+[\.)]|[一二三四五六七八九十]+[、\.)]|\([一二三四五六七八九十]+\))", line
        ):
            list_like_count += 1

    if list_like_count >= max(2, len(lines) // 2):
        return "list"

    if len(cleaned) <= 60 and not cleaned.endswith(("。", "！", "？")):
        return "title"

    return "paragraph"


def _get_heading_level(style_name: str, text: str = "", paragraph=None):
    """判断标题是几级，识别 Word 自带标题、中文标题和一些常见编号标题。"""
    if not style_name and not text:
        return None

    s = (style_name or text or "").strip().lower()

    if s.startswith("heading"):
        parts = s.split()
        if parts and parts[-1].isdigit():
            return int(parts[-1])

    if s.startswith("标题"):
        digits = "".join(ch for ch in s if ch.isdigit())
        if digits:
            return int(digits)

    if re.match(r"^\d+(\.\d+)*[\.、]?\s+\S+", s):
        level = min(3, s.count(".") + 1)
        return max(1, level)

    if re.match(r"^[一二三四五六七八九十]+[、\.)]\s*\S+", s):
        return 1

    if re.match(r"^（[一二三四五六七八九十]+）\s*\S+", s):
        return 2

    if (
        paragraph is not None
        and len(s) <= 30
        and not s.endswith(("。", "！", "？"))
        and " " not in s
    ):
        if any(run.bold for run in paragraph.runs if run.text.strip()):
            return 3

    return None


def _looks_like_heading(paragraph, text: str) -> bool:
    """看看这一行像不像标题，方便判断是不是结构化文档。"""
    if not text:
        return False

    style_name = paragraph.style.name if paragraph.style else ""
    return _get_heading_level(style_name, text, paragraph) is not None


def _iter_block_items(parent):
    """按文档中的真实顺序遍历段落和表格。"""

    if isinstance(parent, _DocxDocumentType):
        parent_elm = parent.element.body
    else:
        parent_elm = parent._tc

    for child in parent_elm.iterchildren():
        if child.tag.endswith("}p"):
            yield Paragraph(child, parent)
        elif child.tag.endswith("}tbl"):
            yield Table(child, parent)


def _is_list_paragraph(paragraph: Paragraph) -> bool:
    """尽量识别项目符号或编号列表。"""
    text = _normalize_text(paragraph.text)
    if not text:
        return False

    if re.match(
        r"^([\-\*•·]|\d+[\.)]|[一二三四五六七八九十]+[、\.)]|\([一二三四五六七八九十]+\))\s*", text
    ):
        return True

    style_name = paragraph.style.name if paragraph.style else ""
    lower_style = style_name.lower()
    if "list" in lower_style or "项目符号" in style_name or "编号" in style_name:
        return True

    return False


def _paragraph_to_markdown(paragraph: Paragraph) -> Tuple[str, str]:
    """把一个段落转换成 Markdown 文本，并返回内容类型。"""
    text = _normalize_text(paragraph.text)
    if not text:
        return "", "paragraph"

    level = _get_heading_level(paragraph.style.name if paragraph.style else "", text, paragraph)
    if level is not None:
        heading_level = min(max(level, 1), 6)
        return f"{'#' * heading_level} {text}", "title"

    if _is_list_paragraph(paragraph):
        text = re.sub(
            r"^([\-\*•·]|\d+[\.)]|[一二三四五六七八九十]+[、\.)]|\([一二三四五六七八九十]+\))\s*",
            "- ",
            text,
        )
        return text, "list"

    if getattr(paragraph, "level", 0) and paragraph.level > 0:
        indent = "  " * min(paragraph.level, 3)
        return f"{indent}- {text}", "list"

    return text, "paragraph"


def _markdown_table_from_row_values(rows: List[List[str]]) -> str:
    """把二维表格转成 Markdown 表格。"""
    normalized_rows = []
    for row in rows:
        cells = []
        for cell in row:
            cell_text = _normalize_text(cell or "").replace("\n", " ")
            cells.append(cell_text)
        if any(cells):
            normalized_rows.append(cells)

    if not normalized_rows:
        return ""

    header = normalized_rows[0]
    body_rows = normalized_rows[1:]
    header_line = "| " + " | ".join(header) + " |"
    separator_line = "|" + "|".join(["---"] * len(header)) + "|"
    body_lines = ["| " + " | ".join(row) + " |" for row in body_rows]
    return "\n".join([header_line, separator_line, *body_lines])


def _table_to_markdown(table: Table) -> str:
    """把 Word 表格转成 Markdown 表格。"""
    rows = []
    for row in table.rows:
        cells = []
        for cell in row.cells:
            cell_text = _normalize_text(cell.text)
            cells.append(cell_text)
        rows.append(cells)

    return _markdown_table_from_row_values(rows)


def _build_metadata(
    path: str, title_stack, content_type: str, section_id: int, chunk_id: Optional[int] = None
):
    """给每个块补上来源、章节和标题路径，后面检索时更容易追踪。"""
    metadata = {
        "source": path,
        "content_type": content_type,
        "section_id": section_id,
        "title_path": " > ".join(title for _, title in title_stack) if title_stack else "ROOT",
    }
    for level, title in title_stack:
        metadata[f"h{level}"] = title
    if chunk_id is not None:
        metadata["chunk_id"] = chunk_id
    return metadata


def _split_docx_into_markdown_documents(path: str) -> List[Document]:
    """按 Word 的真实阅读顺序，把 DOCX 转成 Markdown 文档。"""
    doc = DocxDocument(path)
    documents = []
    title_stack = []
    section_id = 0
    buffer_lines = []

    def flush_buffer():
        nonlocal section_id
        text = _normalize_text("\n".join(line for line in buffer_lines if line.strip()))
        if not text:
            return

        section_id += 1
        metadata = _build_metadata(path, title_stack, "paragraph", section_id)
        documents.append(Document(page_content=text, metadata=metadata))

    for block in _iter_block_items(doc):
        if isinstance(block, Paragraph):
            raw_text = _normalize_text(block.text)
            if not raw_text:
                continue

            md_text, content_type = _paragraph_to_markdown(block)
            if not md_text:
                continue

            level = _get_heading_level(block.style.name if block.style else "", raw_text, block)
            if level is not None:
                flush_buffer()
                buffer_lines.clear()

                title_stack[:] = [(lv, t) for lv, t in title_stack if lv < level]
                title_stack.append((level, raw_text))
                buffer_lines.append(md_text)
                continue

            buffer_lines.append(md_text)
            continue

        if isinstance(block, Table):
            flush_buffer()
            buffer_lines.clear()

            table_md = _table_to_markdown(block)
            if not table_md:
                continue

            section_id += 1
            metadata = _build_metadata(path, title_stack, "table", section_id)
            documents.append(Document(page_content=table_md, metadata=metadata))

    flush_buffer()

    if not documents:
        fallback_text = _normalize_text("")
        if fallback_text:
            documents.append(
                Document(
                    page_content=fallback_text,
                    metadata={
                        "source": path,
                        "content_type": "paragraph",
                        "section_id": 1,
                        "title_path": "ROOT",
                    },
                )
            )

    return documents


def is_structured_docx(path: str):
    """看看这个 docx 有没有明显的标题结构，有的话就适合做章节化 Markdown 输出。"""
    doc = DocxDocument(path)
    heading_count = 0

    for paragraph in doc.paragraphs:
        text = _normalize_text(paragraph.text)
        if not text:
            continue

        level = _get_heading_level(paragraph.style.name if paragraph.style else "", text, paragraph)
        if level is not None:
            heading_count += 1
            if heading_count >= 2:
                return True

    return False


def unstructure_docx(path: str):
    """没有清晰标题结构时，也直接输出 Markdown 文档，而不是切块。"""
    return _split_docx_into_markdown_documents(path)


def structure_docx(path: str):
    """有标题结构时，输出章节化 Markdown 文档。"""
    return _split_docx_into_markdown_documents(path)


def auto_split_docx(path: str):
    """自动判断文档结构，但最终都返回 Markdown 文档列表。"""
    return _split_docx_into_markdown_documents(path)


def load_docx(path: str):
    """给外部流程调用的统一入口，直接返回 Markdown 化后的文档。"""
    return _split_docx_into_markdown_documents(path)
